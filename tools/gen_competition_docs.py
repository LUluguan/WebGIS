# -*- coding: utf-8 -*-
"""gen_competition_docs.py — 生成竞赛交付 Word 文档(五号宋体, A4 标准页)。
输出到 交付文档/ : 安装配置 / 部署说明 / 数据说明 / 作品介绍(≤8页) / 作品设计文档(全版) / 系统概述(≤1页)
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\Competiton\交付文档"
os.makedirs(OUT, exist_ok=True)

SONG = "宋体"
HEI = "黑体"
FIVE = Pt(10.5)  # 五号


def _set_east_asia(run, cn):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), cn)


def set_font(run, cn=SONG, size=FIVE, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = size
    run.font.bold = bold
    _set_east_asia(run, cn)


def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'
    st.font.size = FIVE
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.append(rf)
    rf.set(qn('w:ascii'), 'Times New Roman')
    rf.set(qn('w:hAnsi'), 'Times New Roman')
    rf.set(qn('w:eastAsia'), SONG)
    st.paragraph_format.line_spacing = 1.25
    st.paragraph_format.space_after = Pt(0)
    return doc


def title(doc, text, sub=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(text), cn=HEI, size=Pt(18), bold=True)
    if sub:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(12)
        set_font(p2.add_run(sub), cn=SONG, size=Pt(12))
    return p


def heading(doc, text, level=1):
    sizes = {1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(10.5)}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), cn=HEI if level <= 2 else SONG, size=sizes.get(level, FIVE), bold=level <= 2)
    return p


def para(doc, text, size=FIVE, bold=False, align=None, indent=True, space_before=0, space_after=0, cn=SONG):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Pt(size.pt * 2)
    if align:
        p.alignment = align
    set_font(p.add_run(text), cn=cn, size=size, bold=bold)
    return p


def add_table(doc, rows, cols, header=None, widths=None):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in range(rows):
        for c in range(cols):
            cell = t.cell(r, c)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            if r == 0 and header:
                set_font(cell.paragraphs[0].add_run(header[c]), cn=HEI, size=Pt(9), bold=True)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t


def fill_cell(cell, text, size=Pt(9)):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        r.text = ""
    set_font(p.add_run(text), size=size)


def table_row(t, row_idx, values, size=Pt(9)):
    for c, v in enumerate(values):
        fill_cell(t.cell(row_idx, c), v, size=size)


def save(doc, name):
    path = os.path.join(OUT, name)
    doc.save(path)
    print("生成:", path)


# ============================================================
# 文档 1: 安装及配置文件说明
# ============================================================
def doc1():
    doc = new_doc()
    title(doc, "安装及配置文件说明", "广东降雨洪涝 WebGIS · 2026 易智瑞杯 GIS 竞赛")
    heading(doc, "一、环境要求")
    para(doc, "操作系统：Windows 10/11（推荐）或 Linux；内存建议 8GB 以上。")
    para(doc, "Python：3.10 及以上（本项目开发环境为 Python 3.13）。")
    para(doc, "可选组件：PostgreSQL 18 + PostGIS 3.6（不安装也可运行，服务层自动回退读取本地数据文件）；Docker（用于容器化部署）。")
    heading(doc, "二、依赖安装")
    para(doc, "项目依赖清单见 requirements.txt，包括：fastapi、uvicorn、psycopg2-binary（Web 服务）、numpy、rasterio、tifffile、Pillow（地理/栅格处理）、requests（真实事件卫星数据管线）。深度学习推理需 torch（CPU 版），通过 PyTorch 官方 CPU 源单独安装：")
    para(doc, "pip install -r requirements.txt", indent=False, cn="Consolas")
    para(doc, "pip install torch --index-url https://download.pytorch.org/whl/cpu", indent=False, cn="Consolas")
    para(doc, "上述步骤已封装在 setup.bat 中（自动创建 .venv 虚拟环境并安装全部依赖），首次使用直接双击 setup.bat 即可。")
    heading(doc, "三、配置文件说明")
    add_table(doc, 6, 2, header=["配置项", "说明"])
    rows = [
        ("FLOOD_DB_HOST / PORT / NAME / USER / PASSWORD", "PostGIS 连接参数，见 .env.example；连不上时服务层自动回退读 flood_out/ 本地文件"),
        ("FLOOD_PORT", "服务端口，默认 8001；run.bat 读取该环境变量"),
        (".env.example", "数据库与端口配置模板（环境变量方式，不写入代码）"),
        (".gitignore", "排除大体积原始数据（pre_*.nc、GF-FloodNet、SRTM、precip_tif 等）"),
        (".dockerignore", "Docker 构建时排除的目录（.venv、precip_tif、缓存等）"),
    ]
    for i, r in enumerate(rows, start=1):
        table_row(doc.tables[-1], i, list(r))
    heading(doc, "四、目录结构")
    add_table(doc, 12, 2, header=["目录/文件", "用途"])
    rows = [
        ("app.py", "FastAPI 服务层（场景/淹没/水深/降雨/UNet 推理/真实事件元数据）"),
        ("welcome.html / index.html / dashboard.html / realevent.html / unet.html", "欢迎页 / 主场景(双模式) / 数据大屏 / 真实事件页 / UNet 演示页"),
        ("web/", "本地化前端资源（Cesium 1.95 + ECharts 5.5，不依赖 CDN）"),
        ("realevent_beijiang.py + sat_data.py + unet_apply.py + sar_change.py", "真实事件管线（卫星影像→UNet→水位反演）"),
        ("unet_model.py / train_unet.py / infer_unet.py / eval_unet.py", "UNet 深度学习训练与推理"),
        ("prep_precip.py / prep_return_period.py / fetch_dem.py / bathtub_flood.py / water_level_inversion.py", "离线数据与算法管线"),
        ("flood_out/", "重现期计算结果（水深 tif / 淹没 geojson / scenarios.json）"),
        ("realevent_out/", "真实事件结果（真彩/掩膜/水深 PNG + depth.tif + realevent.json）"),
        ("dem/  unet_out/", "研究区 DEM / 训练好的 UNet 模型"),
        ("setup.bat / run.bat", "Windows 一键安装 / 启动"),
        ("Dockerfile / docker-compose.yml", "Docker 容器化部署"),
    ]
    for i, r in enumerate(rows, start=1):
        table_row(doc.tables[-1], i, list(r))
    save(doc, "01_安装及配置文件说明.docx")


# ============================================================
# 文档 2: 部署说明文档
# ============================================================
def doc2():
    doc = new_doc()
    title(doc, "部署说明文档", "广东降雨洪涝 WebGIS · 含在线访问地址")
    heading(doc, "一、部署方式（三种任选）")
    heading(doc, "（一）Windows 一键脚本（推荐演示）", 3)
    para(doc, "第 1 步：下载并解压项目，或执行 git clone https://github.com/LUluguan/WebGIS.git 克隆仓库。")
    para(doc, "第 2 步：双击 setup.bat，自动创建虚拟环境、安装依赖（含 torch CPU 版）。")
    para(doc, "第 3 步：双击 run.bat，启动服务并自动打开浏览器进入欢迎页。")
    heading(doc, "（二）Docker 容器化部署", 3)
    para(doc, "在装有 Docker 的机器上执行：docker compose up -d，然后访问 http://localhost:8001/。镜像约 2-3GB（含 torch CPU 版）。PostGIS 可不配置，服务层自动回退本地文件。")
    heading(doc, "（三）命令行方式", 3)
    para(doc, "pip install -r requirements.txt 后执行：uvicorn app:app --host 127.0.0.1 --port 8001。")
    heading(doc, "二、数据库（可选）")
    para(doc, "默认连接本地 flood_analysis 库（PostgreSQL + PostGIS）。数据库不可用时服务层自动回退读取 flood_out/ 本地文件，因此不安装数据库也能完整演示。如需入库，按 .env.example 配置环境变量并运行 load_flood_pg.py。")
    heading(doc, "三、在线访问地址")
    para(doc, "本系统未部署至公共云平台，以下为可访问方式：")
    add_table(doc, 4, 2, header=["访问方式", "地址"])
    rows = [
        ("本机演示（部署后）", "http://127.0.0.1:8001/（欢迎页自动进入）"),
        ("源码仓库（克隆即部署）", "https://github.com/LUluguan/WebGIS.git"),
        ("交互式 API 文档", "http://127.0.0.1:8001/docs"),
    ]
    for i, r in enumerate(rows, start=1):
        table_row(doc.tables[-1], i, list(r))
    heading(doc, "四、演示流程")
    para(doc, "欢迎页 → 三维洪涝模拟（主场景，顶部可切换「模拟·珠江新城 / 真实·英德」双模式）→ 数据大屏 → 真实事件独立页。三维场景需联网加载天地图底图；Cesium/ECharts 已本地化，不依赖 jsdelivr CDN。")
    heading(doc, "五、常见问题")
    add_table(doc, 5, 2, header=["问题", "处理"])
    rows = [
        ("三维场景白屏/无底图", "需联网加载天地图底图；确认网络与 TDT_KEY 有效"),
        ("数据大屏降雨图为空", "precip_tif 体积大未随仓库分发，属预期；其余功能不受影响"),
        ("torch 安装失败", "手动执行 pip install torch --index-url https://download.pytorch.org/whl/cpu"),
        ("需要 PostGIS", "配置 .env 环境变量后重启；服务层优先读库、失败自动回退文件"),
    ]
    for i, r in enumerate(rows, start=1):
        table_row(doc.tables[-1], i, list(r))
    save(doc, "02_部署说明文档.docx")


# ============================================================
# 文档 3: 相关数据说明
# ============================================================
def doc3():
    doc = new_doc()
    title(doc, "相关数据说明", "广东降雨洪涝 WebGIS")
    heading(doc, "一、数据清单")
    add_table(doc, 9, 3, header=["数据", "来源", "用途"])
    rows = [
        ("GLO-30 DEM（30m）", "Copernicus Programme（公开）", "研究区地形、水位反演与水深计算"),
        ("SRTM1 DEM（30m）", "NASA/USGS（公开）", "全省 DEM（初筛，洪涝分析弃用）"),
        ("降雨数据 pre_2021~2025", "ERA5/公开气象再分析（公开）", "逐月降雨，拟合重现期"),
        ("GF-FloodNet 样本集", "公开论文数据集（高分卫星）", "UNet 水体提取训练（13,388 对）"),
        ("Sentinel-1 RTC / Sentinel-2", "Copernicus Programme（公开）", "真实洪涝事件（北江 2022-06 英德）影像"),
        ("广州建筑数据（290 栋）", "公开数据", "三维场景建筑建模"),
        ("flood_out/ 计算结果", "本系统算法管线生成", "重现期水深/淹没范围"),
        ("realevent_out/ 真实事件结果", "本系统 UNet 管线生成", "真实事件掩膜与水深"),
    ]
    for i, r in enumerate(rows, start=1):
        table_row(doc.tables[-1], i, list(r))
    heading(doc, "二、数据涉密声明")
    para(doc, "本作品使用的全部数据均来自公开渠道：美国地质调查局（USGS）、哥白尼计划（Copernicus Programme，含 GLO-30 DEM 与 Sentinel-1/2 卫星影像）、公开论文数据集（GF-FloodNet）以及公开的网络底图服务（天地图）。上述数据均为公开、免费且无涉密内容，不涉及国家秘密、个人隐私或商业机密。")
    para(doc, "原始大体积数据（pre_*.nc 约 4.8GB、GF-FloodNet 约 8.3GB、precip_tif 约 116MB）未随源码仓库分发，仅保留算法管线脚本，评审或部署时可按需重新生成。参赛报名相关信息（*.docx）未提交至公开仓库。")
    para(doc, "作品成果仅用于教学与竞赛演示，不构成任何工程依据或决策参考。")
    save(doc, "03_相关数据说明.docx")


# ============================================================
# 文档 4: 作品介绍文档（≤8 页）
# ============================================================
def doc4():
    doc = new_doc()
    title(doc, "基于 WebGIS 的三维城市降雨洪涝可视化", "作品介绍 · 华南农业大学 · 2026 易智瑞杯 GIS 竞赛")
    heading(doc, "一、需求分析")
    para(doc, "华南地区汛期降雨集中，城市内涝频发。传统洪涝灾害表达以平面专题图、统计报表为主，难以直观呈现洪水在城市三维空间中的分布与淹没深度。本作品面向城市洪涝灾害的可视化与决策支持需求，提出基于 WebGIS 的三维城市降雨洪涝可视化系统，实现不同重现期降雨条件下的城市积水深度与淹没范围三维动态表达，并支持基于真实卫星影像的洪涝事件反演，为洪涝灾害风险评估、应急管理与科普演示提供直观、可交互的可视化工具。")
    heading(doc, "二、系统架构说明")
    para(doc, "系统采用「表现层—服务层—数据层—计算层」四层架构，各层开发工具如下：")
    add_table(doc, 5, 3, header=["层次", "功能", "开发工具"])
    rows = [
        ("表现层", "三维场景 / 数据大屏 / 真实事件页 / 欢迎页", "CesiumJS 1.95（三维）、ECharts 5.5（大屏）、HTML/CSS/JavaScript、天地图 WMTS 底图"),
        ("服务层", "REST API：场景、淹没范围、水深图、降雨、UNet 推理、真实事件", "Python + FastAPI + uvicorn"),
        ("数据层", "地理数据存储与查询（可选）", "PostgreSQL 18 + PostGIS 3.6 + psycopg2"),
        ("计算层", "重现期拟合、浴缸法水位反演、UNet 水体提取、水位反演", "Python + numpy + rasterio + tifffile + PyTorch(CPU)"),
    ]
    for i, r in enumerate(rows, start=1):
        table_row(doc.tables[-1], i, list(r))
    para(doc, "前端三维库与图表库（Cesium、ECharts）已本地化部署于 web/ 目录，不依赖外网 CDN，增强演示可靠性。")
    heading(doc, "三、总体设计")
    heading(doc, "（一）功能设计", 3)
    para(doc, "1. 三维洪涝模拟：珠江新城三维场景，切换 2/5/10/50/100 年重现期，查看浴缸法反演的水位与水深色带，建筑按真实地面高程抬升，广州塔专属建模。2. 真实事件反演：北江 2022-06 英德洪水真实卫星影像，经 UNet 提取淹没范围、边界水位反演得到三维水深。3. 数据大屏：KPI 指标、逐月降雨态势、水深分布与预警等级统计。4. 欢迎页与导航：作品总览与三模块入口。")
    heading(doc, "（二）数据库设计", 3)
    para(doc, "PostGIS 数据库 flood_analysis 主要表：precip_2021~2025（逐月降雨栅格）、gz_tower_buildings（290 栋建筑，Polygon 4326，含 name/height/levels）、flood_scenarios（重现期场景参数）、flood_extent（淹没范围多边形）、flood_depth（水深栅格）。数据库不可用时自动回退读取本地文件。")
    heading(doc, "（三）关键技术", 3)
    add_table(doc, 6, 2, header=["关键技术", "说明"])
    rows = [
        ("Gumbel 重现期拟合", "逐像元年最大月雨量拟合 Gumbel 分布，得到 2/5/10/50/100 年重现期雨量"),
        ("浴缸法水位反演", "建筑剔除 + 陆域体积守恒求水面高程 W，水深 = max(0, W − 地形)"),
        ("UNet 水体提取", "5 波段多光谱输入（GF-2 蓝绿红近红外 + GF-3 SAR），语义分割提取水体"),
        ("边界水位反演", "UNet 掩膜边界像元 DEM 中位数反演真实水面高程 W，得到真实水深"),
        ("真实卫星数据管线", "Sentinel-1 RTC + Sentinel-2 经 Planetary Computer 匿名签名窗口读取构建 5 波段"),
    ]
    for i, r in enumerate(rows, start=1):
        table_row(doc.tables[-1], i, list(r))
    heading(doc, "四、作品亮点")
    para(doc, "1. 真实数据驱动：突破纯模拟局限，接入真实卫星洪涝影像，UNet 提取真实淹没范围并反演真实水深，与浴缸法模拟互为印证。2. 双模一体交互：同一三维场景内无缝切换「模拟情景 / 真实事件」，直观对比。3. 端到端自动化：从降雨数据、DEM 到前端可视化全程脚本化管线，可复现、可扩展。4. 轻量可部署：前端资源本地化、无需数据库即可演示，支持 Windows 一键脚本与 Docker 容器化部署。5. 真实事件演示（北江 2022-06 英德）：反演水面高程 27.32m、淹没 17.9km²、最大水深 3.82m，经 SAR 暗像元与水体特征验证。")
    para(doc, "（本作品全部数据来自公开渠道，无涉密内容，结果仅供教学演示。）")
    save(doc, "04_作品介绍文档.docx")


# ============================================================
# 文档 5: 作品设计文档（全版本，仅评委参阅）
# ============================================================
def doc5():
    doc = new_doc()
    title(doc, "作品设计文档（全版本）", "基于 WebGIS 的三维城市降雨洪涝可视化 · 仅评委参阅，不上网刊登")
    heading(doc, "一、项目背景与目标")
    para(doc, "华南地区汛期降雨集中，城市内涝频发，传统二维专题图难以表达洪水在城市三维空间中的分布与深度。本项目构建一套基于 WebGIS 的三维城市降雨洪涝可视化系统：以珠江新城为研究区，融合重现期情景模拟（Gumbel 拟合 + 浴缸法水位反演）与真实洪涝事件反演（Sentinel 卫星影像 + UNet 水体提取 + 边界水位反演），在 Cesium 三维场景中动态呈现积水深度与淹没范围。")
    heading(doc, "二、总体架构")
    add_table(doc, 5, 2, header=["层次", "组成"])
    rows = [
        ("表现层", "CesiumJS 三维场景 / ECharts 数据大屏 / 真实事件页 / 欢迎页 / UNet 演示页"),
        ("服务层", "FastAPI：/api/scenarios、/api/flood_extent、/api/flood_depth_png、/api/monthly_rain、/api/depth_hist、/api/realevent、/api/predict"),
        ("数据层", "PostgreSQL + PostGIS（可选，自动回退本地文件）：precip_2021~2025、gz_tower_buildings、flood_scenarios、flood_extent、flood_depth"),
        ("计算层", "离线管线：prep_precip → prep_return_period(Gumbel) → fetch_dem → bathtub_flood → 导出；真实事件管线：sat_data(下载) → unet_apply → water_level_inversion"),
    ]
    for i, r in enumerate(rows, start=1):
        table_row(doc.tables[-1], i, list(r))
    heading(doc, "三、数据层设计")
    heading(doc, "（一）数据来源", 3)
    para(doc, "DEM：Copernicus GLO-30（30m）；降雨：pre_2021~2025 逐月栅格（0.1mm 单位）；建筑：广州珠江新城 290 栋公开数据；UNet 训练：GF-FloodNet 公开数据集（13,388 对，5 波段，高分二号多光谱 + 高分三号 SAR）；真实事件：Sentinel-1 RTC（辐射地形校正，10m）与 Sentinel-2 L2A。")
    heading(doc, "（二）数据库表结构", 3)
    add_table(doc, 6, 3, header=["表", "内容", "关键字段"])
    rows = [
        ("precip_2021~2025", "逐月降雨栅格", "12 波段、nodata=-32768"),
        ("gz_tower_buildings", "广州建筑", "name / height / levels / Polygon(4326)"),
        ("flood_scenarios", "重现期场景", "return_period_y / rain_mm / water_level_m / flooded_area_km2"),
        ("flood_extent", "淹没范围", "return_period_y / geometry(Polygon,4326)"),
        ("flood_depth", "水深栅格", "5 重现期 × 4 瓦片，中心水深 4.64~5.04m"),
    ]
    for i, r in enumerate(rows, start=1):
        table_row(doc.tables[-1], i, list(r))
    heading(doc, "四、算法设计")
    heading(doc, "（一）重现期降雨（Gumbel 拟合）", 3)
    para(doc, "逐像元统计 2021-2025 年最大月雨量（5 个年极值样本），以矩法（MOM）拟合 Gumbel 分布，得到 2/5/10/50/100 年重现期雨量。研究区重现期月雨量：2 年 379 / 5 年 410 / 10 年 431 / 50 年 476 / 100 年 495 mm。")
    heading(doc, "（二）浴缸法水位反演", 3)
    para(doc, "将研究区视为不透水浴缸，综合径流系数 C=0.50，由重现期雨量求径流深；剔除建筑占地后进行陆域体积守恒，求解水面高程 W；水深 D = max(0, W − DEM)。各重现期水位：2 年 4.64m → 100 年 5.04m，陆地淹没 12.6%→15.7%。")
    heading(doc, "（三）UNet 水体提取", 3)
    para(doc, "UNet 编码器-解码器结构（7.76M 参数，base=32），输入 5 波段多光谱（256×256），输出水体概率图。训练集为 GF-FloodNet（13,388 对，中国/澳大利亚/巴西等 8 地区），CPU 子集训练，验证集总像素 IoU≈0.82。推理时按训练均值/标准差逐波段归一化。")
    heading(doc, "（四）边界水位反演", 3)
    para(doc, "UNet 掩膜边界像元即水位等高线，取边界 DEM 中位数作为水面高程 W，水深 D = max(0, W − DEM)（仅掩膜内）。该方法由真实影像直接反演水位，无需水文假设，与浴缸法形成对照。")
    heading(doc, "五、真实事件模块（北江 2022-06 英德）")
    para(doc, "事件：2022-06 北江特大洪水，英德城区被淹。数据：Sentinel-1 RTC VV（2022-06-26 洪水中，灾前 06-02 作基线）+ Sentinel-2 L2A（2022-06-23，B2/B3/B4/B8）。下载：Microsoft Planetary Computer STAC 检索 + 匿名签名 + rasterio 窗口重投影（EPSG:32649，10m 统一网格）。管线：5 波段堆栈 → UNet（prob>0.5）→ 边界水位反演。结果：水面高程 W=27.32m、淹没 17.9km²、平均水深 1.31m、最大水深 3.82m。验证：掩膜像元具备水体特征（SAR 低回波、近红外低值、地势低 19m），与 SAR 暗像元 IoU≈0.27，可靠性标注为「中」。")
    heading(doc, "六、服务层设计")
    para(doc, "FastAPI 提供 REST 接口（见架构表），统一 8001 端口，静态托管全部前端页面与本地化资源；数据库不可用时自动回退读取 flood_out/、realevent_out/ 本地文件，实现免数据库演示。")
    heading(doc, "七、前端设计")
    heading(doc, "（一）三维主场景（index.html）", 3)
    para(doc, "CesiumJS + 天地图 WMTS 底图；290 栋建筑按真实地面高程抬升着色；广州塔以堆叠圆柱近似双曲面「小蛮腰」；模拟模式按重现期渲染水面与水深色带，真实模式渲染 UNet 反演水深与水位面；顶部按钮无缝切换双模式，相机自动飞行。")
    heading(doc, "（二）数据大屏（dashboard.html）", 3)
    para(doc, "ECharts 实现 KPI 指标、逐月降雨态势、水深分布直方图、预警等级饼图；数据缺失时优雅占位。")
    heading(doc, "（三）真实事件页（realevent.html）", 3)
    para(doc, "四步管线图（影像→UNet 掩膜→水位 W→水深），图层开关（真彩/掩膜/水深/SAR 水），三维水位面，方法对比面板，验证指标展示。")
    heading(doc, "八、部署方案")
    para(doc, "前端资源（Cesium 1.95 + ECharts 5.5）本地化至 web/，无 CDN 依赖。部署三方式：Windows 一键脚本（setup.bat 装依赖 + run.bat 启动并自动开浏览器）、Docker 容器化（docker compose up -d，镜像约 2-3GB）、命令行（uvicorn app:app --port 8001）。测试：9 个自动化测试套件覆盖数据获取、UNet 应用、API 回退等。")
    heading(doc, "九、作品亮点与展望")
    para(doc, "亮点：真实卫星数据驱动的水深反演、模拟与真实双模式一体交互、端到端脚本化管线、轻量可部署。展望：接入逐时降雨与水文模型提升模拟精度、扩充更多真实事件样本、部署至云端提供在线服务。")
    save(doc, "05_作品设计文档_全版本.docx")


# ============================================================
# 文档 6: 系统概述（≤1 页）
# ============================================================
def doc6():
    doc = new_doc()
    title(doc, "基于 WebGIS 的三维城市降雨洪涝可视化", "系统概述")
    heading(doc, "一、背景介绍")
    para(doc, "华南地区汛期降雨集中，城市内涝频发。本作品面向城市洪涝灾害可视化与决策支持需求，构建基于 WebGIS 的三维城市降雨洪涝可视化系统，以广州市珠江新城为研究区，实现重现期情景模拟与真实洪涝事件反演的二维/三维一体化表达。")
    heading(doc, "二、主要功能")
    para(doc, "（1）三维洪涝模拟：珠江新城三维场景，切换 2/5/10/50/100 年重现期，展示浴缸法反演的水位与淹没水深；（2）真实事件反演：北江 2022-06 英德洪水真实卫星影像，经 UNet 水体提取与边界水位反演得到三维水深；（3）数据大屏：KPI、逐月降雨态势、水深分布与预警等级统计；（4）欢迎页与三模块导航。")
    heading(doc, "三、主要特点")
    para(doc, "真实数据驱动，UNet 从卫星影像提取淹没范围并反演真实水深；模拟与真实双模式同场景切换；数据获取、算法处理到可视化全程脚本化，可复现；前端资源本地化、无需数据库即可演示，支持一键脚本与 Docker 部署。数据均来自公开渠道，无涉密内容，结果仅供教学演示。")
    save(doc, "06_系统概述.docx")


if __name__ == "__main__":
    doc1()
    doc2()
    doc3()
    doc4()
    doc5()
    doc6()
    print("全部生成完成 →", OUT)
